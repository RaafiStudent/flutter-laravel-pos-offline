import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- KONFIGURASI ---
OUTPUT_FILENAME = 'Full_Source_Code.docx'

# Ekstensi file yang akan diambil kodenya
TARGET_EXTENSIONS = ['.dart', '.php', '.xml', '.yaml', '.html', '.css', '.js']

# Folder yang WAJIB dilewati (Sampah/Cache/Library Pihak Ketiga)
IGNORE_FOLDERS = {
    '.git', '.idea', '.vscode', '.dart_tool', 'build', 
    'node_modules', 'vendor', 'storage', 'obj', 'bin', 
    'android', 'ios', 'windows', 'linux', 'web', # Opsional: jika ingin skip folder native platform
    '__pycache__'
}
# Catatan: Saya masukkan 'android' dll di ignore agar dokumen tidak penuh sampah XML bawaan. 
# Jika Boss butuh AndroidManifest, hapus 'android' dari list di atas.

def generate_documentation():
    print("🚀 Memulai Generator Dokumentasi...")
    
    # Buat Dokumen Word Baru
    doc = Document()
    
    # Halaman Judul
    title = doc.add_heading('DOKUMENTASI FULL SOURCE CODE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Project: Kasir Pintar (Laravel + Flutter)\nOffline-First Architecture')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    file_count = 0

    # Mulai Scanning Folder (Walk)
    for root, dirs, files in os.walk("."):
        # 1. Filter Folder: Buang folder yang ada di IGNORE_FOLDERS
        # Kita modifikasi list 'dirs' secara in-place agar os.walk tidak masuk ke sana
        dirs[:] = [d for d in dirs if d not in IGNORE_FOLDERS]

        for file in files:
            # 2. Cek Ekstensi File
            if any(file.endswith(ext) for ext in TARGET_EXTENSIONS):
                file_path = os.path.join(root, file)
                
                # Biar path terlihat rapi (contoh: .\backend\app\Models\User.php)
                relative_path = os.path.relpath(file_path, ".")
                
                # Skip script ini sendiri biar gak ikut dicopy
                if "generate_docs.py" in relative_path:
                    continue

                print(f"📄 Processing: {relative_path}")

                try:
                    # 3. Baca Isi File
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        code_content = f.read()

                    # Jika file kosong, skip
                    if not code_content.strip():
                        continue

                    # 4. Tulis Judul File (Path)
                    heading = doc.add_heading(relative_path, level=2)
                    heading.style.font.size = Pt(14)
                    heading.style.font.color.rgb = RGBColor(0, 51, 102) # Biru Gelap

                    # 5. Tulis Isi Kode
                    paragraph = doc.add_paragraph(code_content)
                    
                    # Formatting Kode (Monospace Font agar rapi seperti di VS Code)
                    paragraph.style.font.name = 'Courier New'
                    paragraph.style.font.size = Pt(9)
                    
                    # Tambahkan spasi antar file
                    doc.add_paragraph('_' * 50) 
                    file_count += 1

                except Exception as e:
                    print(f"❌ Gagal membaca {relative_path}: {e}")

    # Simpan File Akhir
    try:
        doc.save(OUTPUT_FILENAME)
        print("\n" + "="*50)
        print(f"✅ SUKSES! Dokumentasi selesai dibuat.")
        print(f"📂 Total File: {file_count}")
        print(f"📄 Lokasi File: {os.path.abspath(OUTPUT_FILENAME)}")
        print("="*50)
    except PermissionError:
        print(f"\n❌ GAGAL MENYIMPAN! Tutup file '{OUTPUT_FILENAME}' jika sedang dibuka di Word, lalu coba lagi.")

if __name__ == "__main__":
    generate_documentation()